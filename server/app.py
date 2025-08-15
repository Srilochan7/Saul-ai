import os
import io
from fastapi import FastAPI, UploadFile, File, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv

# Load environment variables FIRST
load_dotenv()

from models import SummaryResponse
from file_parser import extract_text
from summarizer import DocumentSummarizer
from document_validator import DocumentValidator

# Check API key after loading .env
api_key = os.getenv("GROQ_API_KEY")
if not api_key:
    raise RuntimeError(
        "GROQ_API_KEY not found in environment. "
        "Please create a .env file with your Groq API key: GROQ_API_KEY=gsk_your_key_here"
    )

# Initialize the summarizer and validator with the API key
summarizer = DocumentSummarizer(api_key)
document_validator = DocumentValidator()

app = FastAPI(
    title="AI Legal Summarizer", 
    version="1.0.0",
    description="AI-powered legal document summarization service"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://127.0.0.1:3000",
        "https://saul-ai.vercel.app",
        "https://saul-ai-*.vercel.app",  # Preview deployments
        "https://*.vercel.app"  # All Vercel deployments
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)

@app.get("/")
def root():
    """Root endpoint"""
    return {
        "message": "AI Legal Summarizer API",
        "version": "1.0.0",
        "docs": "/docs",
        "health": "/health"
    }

@app.get("/health")
def health():
    """Health check endpoint"""
    return {
        "status": "ok",
        "groq_api_configured": bool(api_key),
        "model": os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
    }

@app.post("/summarize", response_model=SummaryResponse)
async def summarize(file: UploadFile = File(...)):
    """
    Summarize a legal document and return JSON response
    
    Accepts: PDF, DOCX, or TXT files
    Returns: Structured JSON with summary, key points, critical clauses, and recommendations
    """
    # Validate file
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
    
    # Check file size (10MB limit)
    MAX_FILE_SIZE = 10 * 1024 * 1024
    if file.size and file.size > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File too large (max 10MB)")
    
    try:
        # Read file data
        data = await file.read()
        if not data:
            raise HTTPException(status_code=400, detail="Empty file uploaded")

        # Extract text
        text = extract_text(data, filename=file.filename, content_type=file.content_type)
        if not text or len(text.strip()) < 50:
            raise HTTPException(
                status_code=400, 
                detail="Could not extract sufficient text from document (minimum 50 characters required)"
            )

        print(f"Processing document: {file.filename}, Text length: {len(text)} characters")

        # Validate document content
        is_legal, reason, confidence = document_validator.validate_document_content(text)
        
        if not is_legal:
            supported_types = document_validator.get_supported_document_types()
            raise HTTPException(
                status_code=422, 
                detail={
                    "message": reason,
                    "suggestion": "Please upload a legal document for analysis.",
                    "supported_types": supported_types,
                    "confidence_score": confidence
                }
            )
        
        print(f"Document validation passed - Confidence: {confidence:.2f}")

        # Summarize document
        result = summarizer.summarise_document(text)
        
        # Add validation info to response
        result['validation'] = {
            "is_legal_document": True,
            "confidence_score": confidence,
            "validation_message": reason
        }
        
        # Validate result structure
        if not isinstance(result, dict):
            raise HTTPException(status_code=500, detail="Invalid response format from summarizer")
        
        return SummaryResponse(**result)
    
    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        print(f"Unexpected error in /summarize: {e}")
        raise HTTPException(status_code=500, detail=f"Processing error: {str(e)}")

@app.post("/validate-document")
async def validate_document(file: UploadFile = File(...)):
    """
    Validate if uploaded document appears to be legal-related
    """
    try:
        data = await file.read()
        if not data:
            raise HTTPException(status_code=400, detail="Empty file uploaded")

        text = extract_text(data, filename=file.filename, content_type=file.content_type)
        if not text or len(text.strip()) < 50:
            raise HTTPException(
                status_code=400, 
                detail="Could not extract sufficient text from document"
            )

        is_legal, reason, confidence = document_validator.validate_document_content(text)
        
        return {
            "is_legal_document": is_legal,
            "reason": reason,
            "confidence_score": confidence,
            "supported_types": document_validator.get_supported_document_types() if not is_legal else None,
            "text_length": len(text),
            "filename": file.filename
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Validation error: {str(e)}")

@app.post("/summarize/docx")
async def summarize_docx(file: UploadFile = File(...)):
    """
    Summarize a legal document and return a formatted Word document
    
    Accepts: PDF, DOCX, or TXT files  
    Returns: Downloadable DOCX file with formatted summary
    """
    try:
        # Read and validate file
        data = await file.read()
        if not data:
            raise HTTPException(status_code=400, detail="Empty file uploaded")
        
        # Extract text
        text = extract_text(data, filename=file.filename, content_type=file.content_type)
        if not text or len(text.strip()) < 50:
            raise HTTPException(
                status_code=400, 
                detail="Could not extract sufficient text from document"
            )
        
        print(f"Processing document for DOCX: {file.filename}, Text length: {len(text)} characters")
        
        # Summarize document
        result = summarizer.summarise_document(text)

        # Create Word document
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('Legal Document Summary', 0)
        
        # Add disclaimer
        disclaimer = doc.add_paragraph()
        disclaimer_run = disclaimer.add_run("DISCLAIMER: ")
        disclaimer_run.bold = True
        disclaimer.add_run(
            "This is an AI-generated legal summary. It is not legal advice and should not be "
            "relied upon for legal decisions. Please consult with a qualified attorney for "
            "legal matters."
        )
        
        # Add spacing
        doc.add_paragraph()

        # Key Points Section
        key_points_heading = doc.add_heading("Key Points", level=1)
        key_points = result.get("key_points", [])
        if key_points:
            for point in key_points:
                doc.add_paragraph(point, style="List Bullet")
        else:
            doc.add_paragraph("No key points identified.", style="List Bullet")
        
        doc.add_paragraph()  # Spacing

        # Critical Clauses Section
        critical_clauses_heading = doc.add_heading("Critical Clauses", level=1)
        critical_clauses = result.get("critical_clauses", [])
        
        if critical_clauses:
            for clause in critical_clauses:
                if isinstance(clause, dict):
                    # Add clause title as subheading
                    clause_title = clause.get("title", "Unnamed Clause")
                    doc.add_heading(clause_title, level=2)
                    
                    # Add clause description
                    description = clause.get("description", "No description provided.")
                    doc.add_paragraph(description)
                    doc.add_paragraph()  # Spacing between clauses
                else:
                    # Handle case where clause is just a string (backward compatibility)
                    doc.add_paragraph(str(clause), style="List Bullet")
        else:
            doc.add_paragraph("No critical clauses identified.")

        # Recommendations Section
        recommendations_heading = doc.add_heading("Recommendations", level=1)
        recommendations = result.get("recommendations", [])
        if recommendations:
            for rec in recommendations:
                doc.add_paragraph(rec, style="List Bullet")
        else:
            doc.add_paragraph("No specific recommendations provided.", style="List Bullet")
        
        doc.add_paragraph()  # Spacing

        # Full Summary Section
        summary_heading = doc.add_heading("Full Summary", level=1)
        summary_text = result.get("full_summary", "No summary available.")
        
        # Split into paragraphs if there are line breaks
        paragraphs = summary_text.split('\n\n') if '\n\n' in summary_text else [summary_text]
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())

        # Save document to buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        # Generate filename
        base_name = (file.filename or "document").rsplit(".", 1)[0]
        output_filename = f"{base_name}_summary.docx"
        
        print(f"Generated DOCX summary: {output_filename}")
        
        return StreamingResponse(
            io.BytesIO(buf.getvalue()),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{output_filename}"',
                "Content-Length": str(len(buf.getvalue()))
            }
        )
    
    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        print(f"Error in /summarize/docx: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX generation error: {str(e)}")

@app.get("/test-connection")
async def test_connection():
    """Test Groq API connection"""
    try:
        test_result = summarizer.test_connection()
        return {
            "status": "success", 
            "message": "Groq API connection successful", 
            "response": test_result,
            "model": summarizer.model
        }
    except Exception as e:
        return {
            "status": "error", 
            "error": str(e),
            "message": "Failed to connect to Groq API"
        }

@app.get("/supported-formats")
def supported_formats():
    """Get list of supported file formats"""
    return {
        "formats": [
            {
                "extension": ".pdf",
                "mime_type": "application/pdf",
                "description": "Portable Document Format"
            },
            {
                "extension": ".docx", 
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "description": "Microsoft Word Document"
            },
            {
                "extension": ".txt",
                "mime_type": "text/plain", 
                "description": "Plain Text File"
            }
        ],
        "max_file_size": "10MB",
        "min_content_length": "50 characters"
    }

# Error handlers
@app.exception_handler(ValueError)
async def value_error_handler(request, exc):
    return HTTPException(status_code=400, detail=str(exc))

@app.exception_handler(RuntimeError)
async def runtime_error_handler(request, exc):
    return HTTPException(status_code=500, detail=str(exc))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "app:app", 
        host="127.0.0.1", 
        port=8000, 
        reload=True,
        log_level="info"
    )